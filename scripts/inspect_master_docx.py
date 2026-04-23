"""One-shot inspector: dumps paragraph-by-paragraph structure of the master
resume so we can map LLM-tailorable sections to anchor indices.

Run: python scripts/inspect_master_docx.py
"""
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

DOC_PATH = Path(__file__).resolve().parent.parent / "data" / "master_resume.docx"

doc = Document(DOC_PATH)

print(f"=== Document sections: {len(doc.sections)} ===")
for i, sec in enumerate(doc.sections):
    print(f"  Section {i}: page={sec.page_width}x{sec.page_height}, "
          f"margins L/R/T/B = {sec.left_margin}/{sec.right_margin}/"
          f"{sec.top_margin}/{sec.bottom_margin}")

print(f"\n=== Tables: {len(doc.tables)} ===")
for i, tbl in enumerate(doc.tables):
    print(f"  Table {i}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            text_preview = (cell.text[:60] + "…") if len(cell.text) > 60 else cell.text
            print(f"    [{r}][{c}]: {text_preview!r}")

print(f"\n=== Paragraphs: {len(doc.paragraphs)} ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "?"
    text_preview = (p.text[:90] + "…") if len(p.text) > 90 else p.text
    run_info = []
    for r in p.runs:
        font = r.font
        flags = []
        if r.bold: flags.append("B")
        if r.italic: flags.append("I")
        size = font.size.pt if font.size else None
        name = font.name or "?"
        preview = (r.text[:25] + "…") if len(r.text) > 25 else r.text
        run_info.append(f"[{'/'.join(flags) or '-'} {name} {size}pt] {preview!r}")
    print(f"  P{i:03d} [{style}] {text_preview!r}")
    for ri in run_info:
        print(f"        {ri}")
