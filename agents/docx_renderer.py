"""Render tailored resume bytes by replacing [[TAGS]] in the master_resume.docx.

Strategy: open the master DOCX, scan all paragraphs for specific tags (e.g., [[ROLE_1_BULLET_1]]),
and replace them with tailored content. Paragraph-level formatting (bullet glyphs,
indentation, tab stops, fonts) is preserved because we only mutate runs, never pPr.

Bullet text from the LLM carries ** markers for inline emphasis. The parser
splits on ** and emits alternating bold/normal runs whose font inherits from
the master's Garamond 12pt body style.
"""
from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from agents.agent_1 import TailoredAssets

logger = logging.getLogger(__name__)

# Body font used throughout — matches master's run-level formatting.
_BODY_FONT_NAME = "Garamond"
_BODY_FONT_SIZE_PT = 12

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _qn(tag: str) -> str:
    """Qualify a w:-prefixed tag name with the wordprocessingml namespace."""
    return _W_NS + tag


def _parse_bold_segments(text: str) -> list[tuple[str, bool]]:
    """Split LLM bullet text on ** markers into (text, is_bold) segments."""
    if not text:
        return []
    if "**" not in text:
        return [(text, False)]

    parts = text.split("**")
    if len(parts) % 2 == 0:
        logger.warning(f"Unbalanced ** in bullet, rendering as plain text: {text[:60]!r}")
        return [(text, False)]

    segments: list[tuple[str, bool]] = []
    for i, part in enumerate(parts):
        if not part:
            continue
        segments.append((part, i % 2 == 1))
    return segments


def _apply_body_font(run) -> None:
    """Set Garamond 12pt on a run (idempotent, matches master body style)."""
    run.font.name = _BODY_FONT_NAME
    run.font.size = Pt(_BODY_FONT_SIZE_PT)


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    """Remove all <w:r> children from the paragraph while keeping pPr intact."""
    p_elem = paragraph._p
    for r in list(p_elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")):
        p_elem.remove(r)


def _rewrite_paragraph(paragraph: Paragraph, text: str) -> None:
    """Clear a paragraph's runs and emit new ones from ** bold markers in text."""
    _clear_paragraph_runs(paragraph)
    for segment_text, is_bold in _parse_bold_segments(text):
        run = paragraph.add_run(segment_text)
        run.bold = is_bold
        _apply_body_font(run)


def _rewrite_skill_line(paragraph: Paragraph, category: str, skills: list[str]) -> None:
    """Emit '<b>Category:</b> skill1, skill2, ...' preserving paragraph formatting."""
    _clear_paragraph_runs(paragraph)

    label = paragraph.add_run(f"{category}:")
    label.bold = True
    _apply_body_font(label)

    body = paragraph.add_run(f" {', '.join(skills)}.")
    body.bold = False
    _apply_body_font(body)


def _rewrite_interests_line(paragraph: Paragraph, interests: str) -> None:
    """Emit '<b>Interests:</b> ...' preserving paragraph formatting."""
    _clear_paragraph_runs(paragraph)

    label = paragraph.add_run("Interests:")
    label.bold = True
    _apply_body_font(label)

    body_text = interests.strip()
    if not body_text.endswith("."):
        body_text += "."
    body = paragraph.add_run(f" {body_text}")
    body.bold = False
    _apply_body_font(body)


def render_tailored_docx(assets: TailoredAssets, master_path: Path) -> bytes:
    """Return tailored DOCX bytes by replacing [[TAGS]] in the master's paragraphs.
    
    Tags supported:
    - [[ROLE_1_BULLET_N]] (N=1..5)
    - [[ROLE_2_BULLET_N]] (N=1..3)
    - [[EDUCATION_BULLET_N]] (N=1..2)
    - [[SKILLS_Languages & Backend]], [[SKILLS_Frontend & Architecture]], etc.
    - [[INTERESTS]]
    """
    doc = Document(master_path)
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
            
        # 1. Match Experience Role 1 Bullets
        for i in range(1, 6):
            tag = f"[[ROLE_1_BULLET_{i}]]"
            if tag in text:
                bullets = assets.experience[0].bullets if len(assets.experience) > 0 else []
                val = bullets[i-1] if i <= len(bullets) else ""
                _rewrite_paragraph(paragraph, val)
                break
        else:
            # 2. Match Experience Role 2 Bullets
            for i in range(1, 4):
                tag = f"[[ROLE_2_BULLET_{i}]]"
                if tag in text:
                    bullets = assets.experience[1].bullets if len(assets.experience) > 1 else []
                    val = bullets[i-1] if i <= len(bullets) else ""
                    _rewrite_paragraph(paragraph, val)
                    break
            else:
                # 3. Match Education Bullets
                for i in range(1, 3):
                    tag = f"[[EDUCATION_BULLET_{i}]]"
                    if tag in text:
                        bullets = assets.education.bullets if assets.education else []
                        val = bullets[i-1] if i <= len(bullets) else ""
                        _rewrite_paragraph(paragraph, val)
                        break
                else:
                    # 4. Match Skills
                    for category in assets.skills.keys():
                        tag = f"[[SKILLS_{category}]]"
                        if tag in text:
                            _rewrite_skill_line(paragraph, category, assets.skills[category])
                            break
                    else:
                        # 5. Match Interests
                        if "[[INTERESTS]]" in text:
                            _rewrite_interests_line(paragraph, assets.interests or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
